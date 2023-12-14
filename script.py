# from docx import Document

# def docx_to_txt(docx_path, txt_path):
#     doc = Document(docx_path)
#     text_content = ""
#     for paragraph in doc.paragraphs:
#         text_content += paragraph.text + "\n"

#     with open(txt_path, "w", encoding="utf-8") as txt_file:
#         txt_file.write(text_content)

# if __name__ == "__main__":
#     # Replace 'input.docx' and 'output.txt' with your actual file paths
#     docx_file_path = "./16.docx"
#     txt_file_path = "./text_output.txt"

#     docx_to_txt(docx_file_path, txt_file_path)

import docx
from docx import Document
from docx.table import Table

def extract_flagged_entries(docx_path, txt_path):
    doc = Document(docx_path)
    flagged_entries = []
    # Check for flags in paragraphs based on formatting (e.g., bold)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if is_flagged(run): # Implement logic to check for flag based on formatting
                flagged_entries.append(run.text)

    # Check for flags in tables and lists
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if is_flagged(run):
                            flagged_entries.append(run.text)

    # for list_item in doc.element.list_items:
    #     for paragraph in list_item.paragraphs:
    #         for run in paragraph.runs:
    #             if is_flagged(run):
    #                 flagged_entries.append(run.text)

    with open(txt_path, "w", encoding="utf-8") as txt_file:
        txt_file.write("\n".join(flagged_entries))

def is_flagged(run):
    text = run.text
    # Check if the run contains an image (InlineShape)
    if run.element is not None and run.element.tag.endswith('blip'):
        print(text)
        return True
    return False
    # Implement logic to check for flag based on font style, color, etc.
    # return run.bold or run.font.color.name == "yellow"

if __name__ == "__main__":
    # Replace file paths
    docx_path = "./1.docx"
    txt_path = "./flagged_entries.txt"
    extract_flagged_entries(docx_path, txt_path)
