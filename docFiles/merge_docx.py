# from docx import Document

# files = [
#     "1.docx",
#     "10.docx",
#     "11.docx",
#     "12.docx",
#     "13.docx",
#     "14.docx",
#     "15.docx",
#     "16.docx",
#     "17.docx",
#     "18.docx",
#     "2.docx",
#     "3.docx",
#     "4.docx",
#     "5.docx",
#     "6.docx",
#     "7.docx",
#     "8.docx",
#     "9.docx",
# ]

# def combine_word_documents(files):
#     merged_document = Document()

#     for index, file in enumerate(files):
#         sub_doc = Document(file)

#         # Don't add a page break if you've reached the last file.
#         if index < len(files)-1:
#            sub_doc.add_page_break()

#         for element in sub_doc.element.body:
#             merged_document.element.body.append(element)

#     merged_document.save('merged.docx')

# combine_word_documents(files)


from docx import Document
import os

folder_path = './'

# Loop through all files in the folder
files = os.listdir(folder_path)

# Filter out files with the name '.DS_Store'
files = [file for file in files if file != '.DS_Store']
files = [file for file in files if file != 'merge_docx.py']

print("*********************")
print(files)

def combine_word_documents(files, output_file="merged.docx"):
    merged_document = Document()

    for file in files:
        sub_doc = Document(file)

        # Copy paragraphs and styling from sub_doc to merged_document
        for element in sub_doc.element.body:
            merged_document.element.body.append(element)

        # Add a page break between documents
        merged_document.add_page_break()

    # Remove the last page break
    if len(merged_document.paragraphs) > 0 and merged_document.paragraphs[-1].runs == []:
        merged_document.paragraphs.pop()

    merged_document.save(output_file)

combine_word_documents(files)
