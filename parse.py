from docx import Document
import csv
import os
import re

finalCSV = []

def read_docx_table(file_path, state_name = ""):
    document = Document(file_path)
    # print(len(document.tables)) # DEBUG

    tempRowOfCSV = []

    for table in document.tables:
        
        # print("|||||||||| NEW TABLE |||||||||||||||||||| NEW TABLE |||||||||||||||||||||| NEW TABLE |||||||||||") # DEBUG

        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text)

            print(row_data) # DEBUG

            if row_data[0] == "Document Information:": # ['Document Information:', 'Supreme Court of Alabama. January 11, 1910 164 Ala. 111 51 So. 424\nExtracted from page: 1\n']
                tempRowOfCSV = []
                # Document Information
                if row_data[1]:
                    tempRowOfCSV.append(row_data[1])

            if row_data[0] == "WestCheck Information:": # ['WestCheck Information:', 'Birmingham Ry., Light & Power Co. v. Moseley, 164 Ala. 111, 51 So. 424 (Ala. Jan. 11, 1910)\n']
                # WestCheck Information
                if row_data[1]:
                    tempRowOfCSV.append(row_data[1])
                    
                    # CASE ID DOES NOT EXIST FOR ALL CASES PRIMARY CASE INFO
                    # tempRowOfCSV.append(row_data[1].split(",")[2])

                    # Logic To Extract the Primary Case Date.
                    # TODO: LOGIC IS FLAWED FOR SOME ENTRIES WHERE THERE ARE MULTIPLE CIRCULAR BRACKETS
                    pattern = r'\((.*?)\)'
                    matches = re.findall(pattern, row_data[1])
                    # Remove the state from date
                    if matches:
                        tempRowOfCSV.append(' '.join(matches[0].split()[1:]))
                        # TODO: STATE LOGIC IS FLAWED, JUST USED EXTERNALLY PASSED TO THE FUNCTION VARIABLE
                        tempRowOfCSV.append(matches[0].split(' ', 1)[0])
                    else:
                        tempRowOfCSV.append("")
                        tempRowOfCSV.append("")
                    

            # ['Treatment', 'Title', 'Date', 'Type', 'Depth', 'Headnote(s)']
            # Treatment - Overruled by
            if row_data[0] == "Overruled by": # ['Criticized in', ' 1.  Bradley v. Deaton  \n94 So. 767 , 208 Ala. 582 , Ala. , (NO. 6 DIV. 460 )\n', 'Dec. 14, 1922', 'Case', '', 'So.\n']
                # print("@@@@@@@@@@@@@@@@") # DEBUG
                localCopyOfTempRowOfCSV = tempRowOfCSV[:]

                localCopyOfTempRowOfCSV.append(row_data[0])

                # Title
                if row_data[1]:
                    localCopyOfTempRowOfCSV.append(row_data[1])
                # Date
                if row_data[2]:
                    localCopyOfTempRowOfCSV.append(row_data[2])
                
                finalCSV.append(localCopyOfTempRowOfCSV[:])
                
            
            if row_data[0] == "Abrogated by":
                # print("&&&&&&&&&&&&&&&&") # DEBUG
                localCopyOfTempRowOfCSV = tempRowOfCSV[:]

                localCopyOfTempRowOfCSV.append(row_data[0])

                # Title
                if row_data[1]:
                    localCopyOfTempRowOfCSV.append(row_data[1])
                # Date
                if row_data[2]:
                    localCopyOfTempRowOfCSV.append(row_data[2])

                finalCSV.append(localCopyOfTempRowOfCSV[:])

                
            if row_data[0] == "Disavowed by":
                # print("*****************") # DEBUG
                localCopyOfTempRowOfCSV = tempRowOfCSV[:]

                localCopyOfTempRowOfCSV.append(row_data[0])

                # Title
                if row_data[1]:
                    localCopyOfTempRowOfCSV.append(row_data[1])
                # Date
                if row_data[2]:
                    localCopyOfTempRowOfCSV.append(row_data[2])
                
                finalCSV.append(localCopyOfTempRowOfCSV[:])
        
# Define the folder containing the DOCX files for a perticular State
folder_path = '/Users/kdhyani/desktop/west-law/files/Nebraska/Process' #'/Users/kdhyani/desktop/west-law/files/<STATE>/Process'

# print(os.listdir(folder_path)) #DEBUG

# Iterate over all files in the folder
for file_name in os.listdir(folder_path):
    print(file_name)
    if file_name.endswith('.docx'):
        # Construct the absolute path to the file
        docx_file_path = os.path.join(folder_path, file_name)
        read_docx_table(docx_file_path, "Nebraska")

final_output_file_name = "./output/" + "Nebraska" + ".csv"

with open(final_output_file_name, 'w') as f:
    
    # using csv.writer method from CSV package
    write = csv.writer(f)
    
    write.writerow(["DOCUMENT INFORMATION", "PRIMARY CASE INFORMATION", "PRIMARY CASE DATE", "STATE", "TREATMENT", "TREATED CASE INFORMATION", "TREATED CASE DATE"])
    write.writerows(finalCSV)

# for row in table_data:
#     print(row)